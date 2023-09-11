import sys
import os
import fitz  # PyMuPDF
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QWidget, QVBoxLayout, QPushButton, QSplitter, QTextEdit, QLabel, QScrollArea
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap, QImage, QTextCursor, QTextCharFormat, QColor
import xml.etree.ElementTree as ET
from translate import Translator
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
from pdf2docx import Converter
import os
import zipfile
from io import BytesIO
from PIL import Image
import itertools
from docx import Document
from docx.shared import Pt, RGBColor


# Function to extract and save images from a DOCX file
def save_images_from_docx(docx_file, output_folder):
    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Open the DOCX file as a Zip archive
    with zipfile.ZipFile(docx_file, 'r') as docx_zip:
        # Find and extract all image files from the archive
        image_files = [name for name in docx_zip.namelist() if name.startswith('word/media/')]
        for image_file in image_files:
            with docx_zip.open(image_file) as image_data:
                # Load the image using PIL
                image = Image.open(image_data)

                # Generate a unique filename and save the image to the output folder
                image_filename = os.path.basename(image_file)
                image_path = os.path.join(output_folder, image_filename)
                image.save(image_path)
                
def pdf_to_docx(pdf_file, docx_file):
    cv = Converter(pdf_file)
    cv.convert(docx_file, start=0, end=None)
    cv.close()


def style_to_attributes(paragraph, element):
    # Function to extract and add style attributes to the XML element

    # Check if there are runs in the paragraph
    if paragraph.runs:
        # Extract font properties from the first run
        font = paragraph.runs[0].font

        # Extract and add font properties if available
        if font.name:
            element.set("font_name", font.name)
        if font.size:
            element.set("font_size", str(font.size.pt))
        if font.bold:
            element.set("font_bold", "true")
        if font.italic:
            element.set("font_italic", "true")
        if font.underline:
            element.set("font_underline", "true")
        if font.color.rgb:
            element.set("font_color", str(font.color.rgb))

    # Extract alignment
    alignment = paragraph.alignment
    if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
        element.set("alignment", "left")
    elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        element.set("alignment", "center")
    elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
        element.set("alignment", "right")
    elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        element.set("alignment", "justify")

    # Extract spacing properties
    spacing = paragraph.paragraph_format
    if spacing.space_before:
        element.set("space_before", str(spacing.space_before.pt))
    if spacing.space_after:
        element.set("space_after", str(spacing.space_after.pt))
    if spacing.line_spacing:
        element.set("line_spacing", str(spacing.line_spacing))

def note_text_from_table(table, root):
    # Function to extract and note text from a table consisting of one cell

    # Check if the table has only one cell
    if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
        cell = table.rows[0].cells[0]
        
        # Create the <term_table> element
        term_table_element = ET.Element("term_table")
        root.append(term_table_element)
        
        # Counter for naming <text> tags within the <term_table>
        text_counter = 1
        
        for paragraph in cell.paragraphs:
            element = ET.Element(f"text_term_table_{text_counter}")
            text_counter += 1
            style_to_attributes(paragraph, element)
            element.text = paragraph.text
            term_table_element.append(element)

def style_to_attributes(paragraph, element):
    run = paragraph.runs[0] if paragraph.runs else None

    # Font style
    if run and run.font.name:
        element.set("font", run.font.name)
    else:
        element.set("font", "Arial")  # Default font if not specified

    # Font size
    if run and run.font.size:
        element.set("font_size", str(run.font.size.pt))
    else:
        element.set("font_size", "12")  # Default font size if not specified

    # Bold, italic, and underline
    if run:
        element.set("bold", "true" if run.bold else "false")
        element.set("italic", "true" if run.italic else "false")
        element.set("underline", "true" if run.underline else "false")
    else:
        element.set("bold", "false")
        element.set("italic", "false")
        element.set("underline", "false")

    # Typographic alignment
    if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
        element.set("alignment", "left")
    elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        element.set("alignment", "center")
    elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
        element.set("alignment", "right")
    elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        element.set("alignment", "justify")
    else:
        element.set("alignment", "unknown")

    # Leading (line spacing)
    if paragraph.paragraph_format.line_spacing is not None:
        # Convert line spacing value from lines to points (1 line = 12 points)
        leading_value = paragraph.paragraph_format.line_spacing * 12
        element.set("leading", str(leading_value))
    else:
        element.set("leading", "normal")  # Default leading if not specified

def docx_to_xml(docx_file, _xml_file_):
    doc = Document(docx_file)
    root = ET.Element("document")
    is_table = False
    line_number = 1  # Initialize line number counter
    
    # Initialize position variables
    current_left = 0
    current_top = 0

    # Initialize variables to track headers and footers
    header_text = ""
    footer_text = ""

    for paragraph in doc.paragraphs:
        element = ET.Element("text")

        # Handle headings
        if paragraph.style.name.startswith("Heading"):
            element.tag = "header"

        # Handle numbered lists
        if re.match(r'^\d+\.', paragraph.text):
            element.tag = "list_item"

        # Handle links to sources (assuming URLs start with "http" or "www")
        if re.search(r'http|www', paragraph.text):
            element.tag = "link"

        # Set style attributes including typographic alignment and leading
        style_to_attributes(paragraph, element)

        # Extract and add the text content
        element.text = paragraph.text

        # Capture and add the indentation (in points) as an attribute if defined
        if paragraph.paragraph_format.left_indent is not None:
            element.set("indent", str(paragraph.paragraph_format.left_indent.pt))

        # Check if we're inside a table
        if is_table:
            # Add the element to the last cell of the current row
            current_row[-1].append(element)
        else:
            # Add the element to the root
            root.append(element)
        
        # Count the line number for the current element and add it as an attribute
        element.set("line_number", str(line_number))
        line_number += 1  # Increment line number
        
        # Check for headers and footers
        if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            header_text += paragraph.text + "\n"
        elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            footer_text += paragraph.text + "\n"

    # Handle tables
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
            # Note text from tables consisting of one cell
            note_text_from_table(table, root)
        else:
            # Add the entire table to the root
            table_element = ET.Element("table")
            for row in table.rows:
                current_row = []
                row_element = ET.SubElement(table_element, "row")
                for cell in row.cells:
                    cell_element = ET.SubElement(row_element, "cell")
                    current_row.append(cell_element)
                    for paragraph in cell.paragraphs:
                        element = ET.Element("text")

                        # Set style attributes for text in table cells
                        style_to_attributes(paragraph, element)

                        # Extract and add the text content
                        element.text = paragraph.text

                        # Add the element to the cell
                        cell_element.append(element)

            # Add the table element to the root
            root.append(table_element)

    # Handle pictures
    for shape in doc.inline_shapes:
        if shape.type == 3:  # Check if it's a picture
            picture_element = ET.Element("picture")

            # Extract and add the size attributes
            picture_element.set("width", str(shape.width.pt))
            picture_element.set("height", str(shape.height.pt))

            # Calculate and add the position attributes
            picture_element.set("left", str(current_left))
            picture_element.set("top", str(current_top))

            # Add the picture element to the root
            root.append(picture_element)

        # Update the position variables for the next inline shape
        current_left += shape.width.pt
        current_top += shape.height.pt

    # Note headers and footers in the XML structure
    if header_text:
        header_element = ET.Element("header")
        header_element.text = header_text
        root.append(header_element)

    if footer_text:
        footer_element = ET.Element("footer")
        footer_element.text = footer_text
        root.append(footer_element)

    # Save the XML file
    tree = ET.ElementTree(root)
    tree.write(_xml_file_, encoding="utf-8", xml_declaration=True)

def defaultXML2XML(default_xml, xml_structure):
    #required funstions
    def remove_year_and_em_dash(sentence):
        # Define a regular expression pattern to match the year (four digits) and surrounding em dash and spaces
        pattern = r'—?\s?\d{4}\b'

        # Use re.sub to replace the matched pattern with an empty string
        result_sentence = re.sub(pattern, '', sentence)

        # Remove any extra whitespace that may have resulted from the removal
        result_sentence = ' '.join(result_sentence.split())

        return result_sentence

    def convert_russian_date(input_date):
                # Define the mapping of Russian month names to their numeric values
                months = {
                    'января': 1,
                    'февраля': 2,
                    'марта': 3,
                    'апреля': 4,
                    'мая': 5,
                    'июня': 6,
                    'июля': 7,
                    'августа': 8,
                    'сентября': 9,
                    'октября': 10,
                    'ноября': 11,
                    'декабря': 12
                }

                # Split the Russian date string and extract day, month, and year
                date_parts = input_date.split()
                day = int(date_parts[0])
                month = months[date_parts[1]]
                year = int(date_parts[2])

                # Create a datetime object
                dt = datetime(year, month, day)

                # Format the datetime object in the desired output format
                formatted_date = dt.strftime('%Y—%m—%d')

                return formatted_date

    # Function to check if a string contains at least one uppercase word
    def contains_uppercase_word(text):
        words = text.split()
        for word in words:
            if any(char.isupper() for char in word):
                return True
        return False

    def find_largest_line_number(xml_file, **attributes):
        # Initialize the largest line number
        largest_line_number = -1

        # Load the XML file
        tree = ET.parse(xml_file)
        root = tree.getroot()

        # Find all <text> elements with the specified attributes (if provided)
        for text_element in root.findall('.//text'):
            if all(text_element.get(attr) == value for attr, value in attributes.items()):
                if 'line_number' in text_element.attrib:
                    line_number = int(text_element.get('line_number'))
                    largest_line_number = max(largest_line_number, line_number)

        return largest_line_number

    # Load the original XML file
    xml_file = default_xml#f"{base_filename}/{base_filename}_default.xml"#'43.4.27-2020_default.xml'
    tree = ET.parse(xml_file)
    root = tree.getroot()

    # Create a translator instance
    translator = Translator(from_lang='ru', to_lang='en')

    # Create a new XML structure
    standard_root = ET.Element('standard', attrib={
        'xmlns:tbx': 'urn:iso:std:iso:30042:ed-1',
        'xmlns:svg': 'http://www.w3.org/2000/svg'
    })

    front = ET.SubElement(standard_root, 'front')
    iso_meta = ET.SubElement(front, 'iso-meta')
    title_wrap_ru = ET.SubElement(iso_meta, 'title-wrap', {'xml:lang': 'ru'})
    intro_ru = ET.SubElement(title_wrap_ru, 'intro')
    main_ru = ET.SubElement(title_wrap_ru, 'main')
    compl_ru = ET.SubElement(title_wrap_ru, 'compl')
    full_ru = ET.SubElement(title_wrap_ru, 'full')

    title_wrap_eng = ET.SubElement(iso_meta, 'title-wrap', {'xml:lang': 'eng'})
    intro_eng = ET.SubElement(title_wrap_eng, 'intro')
    main_eng = ET.SubElement(title_wrap_eng, 'main')
    compl_eng = ET.SubElement(title_wrap_eng, 'compl')
    full_eng = ET.SubElement(title_wrap_eng, 'full')

    # Add the <doc-ident> structure within <iso-meta>
    doc_ident = ET.SubElement(iso_meta, 'doc-ident')
    sdo = ET.SubElement(doc_ident, 'sdo')
    proj_id_sdo = ET.SubElement(doc_ident, 'proj-id')
    language = ET.SubElement(doc_ident, 'language')
    language.text = 'ru'  # Set the language to 'ru'
    release_version_sdo = ET.SubElement(doc_ident, 'release-version')

    # Add the <std-ident> structure within <iso-meta>
    std_ident = ET.SubElement(iso_meta, 'std-ident')
    originator = ET.SubElement(std_ident, 'originator')
    doc_type = ET.SubElement(std_ident, 'doc-type')
    doc_number = ET.SubElement(std_ident, 'doc-number')
    edition = ET.SubElement(std_ident, 'edition')
    version = ET.SubElement(std_ident, 'version')

    # Add the <std-ref type="dated"> structure within <iso-meta>
    std_ref_dated = ET.SubElement(iso_meta, 'std-ref', {'type': 'dated'})
    std_ref_short = ET.SubElement(iso_meta, 'std-ref', {'type': 'short'})
    std_ref_undated = ET.SubElement(iso_meta, 'std-ref', {'type': 'undated'})
    iso_meta_doc_ref = ET.SubElement(iso_meta, 'doc-ref')
    iso_meta_comm_ref = ET.SubElement(iso_meta, 'comm-ref')
    
    # Add the <secretariat> tag within <iso-meta>
    secretariat = ET.SubElement(iso_meta, 'secretariat')

    # Add the <ics> structure within <iso-meta>
    ics = ET.SubElement(iso_meta, 'ics')

    # Add the <permissions> structure with <copyright-statement>
    permissions = ET.SubElement(iso_meta, 'permissions')
    copyright_statement = ET.SubElement(permissions, 'copyright-statement')
    copyright_holder = ET.SubElement(permissions, 'copyright_holder')
    copyright_year = ET.SubElement(permissions, 'copyright-year')

    # Define a regex pattern to match the year and sentence
    pattern = r'\b\d{4}\b'

    # Add the <custom-meta-group> structure with <custom-meta>
    custom_meta_group = ET.SubElement(iso_meta, 'custom-meta-group')

    custom_meta_udk = ET.SubElement(custom_meta_group, 'custom-meta')
    meta_name_udk = ET.SubElement(custom_meta_udk, 'meta-name')
    meta_value_udk = ET.SubElement(custom_meta_udk, 'meta-value')

    custom_meta_keywords = ET.SubElement(custom_meta_group, 'custom-meta')
    meta_name_keywords = ET.SubElement(custom_meta_keywords, 'meta-name')
    meta_value_keywords = ET.SubElement(custom_meta_keywords, 'meta-value')

    custom_meta_type = ET.SubElement(custom_meta_group, 'custom-meta')
    meta_name_type = ET.SubElement(custom_meta_type, 'meta-name')
    meta_value_type = ET.SubElement(custom_meta_type, 'meta-value')

    meta_name_type.text = 'Type'
    type_meta_value = 'НАЦИОНАЛЬНЫЙ СТАНДАРТ РОССИЙСКОЙ ФЕДЕРАЦИИ'
    meta_value_type.text = type_meta_value 

    nat_meta = ET.SubElement(front, 'nat-meta')

    # Add the <std-ident> structure within <nat-meta>
    std_ident_nat_meta = ET.SubElement(nat_meta, 'std-ident')
    originator_nat_meta = ET.SubElement(std_ident_nat_meta, 'originator')
    doc_type_nat_meta = ET.SubElement(std_ident_nat_meta, 'doc-type')
    doc_number_nat_meta = ET.SubElement(std_ident_nat_meta, 'doc-number')
    edition_nat_meta = ET.SubElement(std_ident_nat_meta, 'edition')
    versionnat_meta_nat_meta = ET.SubElement(std_ident_nat_meta, 'version')
    meta_date_nat_meta_ratification = ET.SubElement(std_ident_nat_meta, 'meta-date', {'type': 'ratification'})
    meta_date_nat_meta_implementation = ET.SubElement(std_ident_nat_meta, 'meta-date', {'type': 'implementation'})
    meta_date_nat_meta_termination = ET.SubElement(std_ident_nat_meta, 'meta-date', {'type': 'termination'})
    comm_ref_nat_meta = ET.SubElement(std_ident_nat_meta, 'comm-ref')
    secretariat_nat_meta = ET.SubElement(std_ident_nat_meta, 'secretariat')

    # Create the <permissions> structure
    permissions_nat_meta = ET.SubElement(nat_meta,'permissions')
    copyright_statement_nat_meta = ET.SubElement(permissions_nat_meta, 'copyright-statement')
    copyright_statement_nat_meta.text = 'Все права сохраняются'
    copyright_year_nat_meta = ET.SubElement(permissions_nat_meta, 'copyright-year')
    copyright_holder_nat_meta = ET.SubElement(permissions_nat_meta, 'copyright-holder')

    notes_front = ET.SubElement(front, 'notes')
    p_notes_front = ET.SubElement(notes_front, 'p')

    sec_foreword = ET.SubElement(front, 'sec', {'id': 'sec_foreword_1', 'sec-type': 'foreword'})
    title_foreword = ET.SubElement(sec_foreword, 'title')
    title_foreword.text = 'Предисловие'
    p1_sec_foreword = ET.SubElement(sec_foreword, 'p')
    p2_sec_foreword = ET.SubElement(sec_foreword, 'p')
    p3_sec_foreword = ET.SubElement(sec_foreword, 'p')
    p4_sec_foreword = ET.SubElement(sec_foreword, 'p')

    sec_intro = ET.SubElement(front, 'sec', {'id': 'sec_intro', 'sec-type': 'intro'})
    title = ET.SubElement(sec_intro, 'title')
    sec_intro_p = ET.SubElement(sec_intro, 'p')

    # <body> element
    body = ET.SubElement(standard_root, 'body')

    # Create a <back> tag and add it to the root
    back = ET.SubElement(standard_root, 'back')

    app_line = None
    bib_line = None
    # Iterate through the <text> tags in the original XML
    for text_element in root.findall('.//text'):
        # Check if the attributes match the specified criteria for <intro>
        if (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '18.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '67.8'
        ):
            # Extract the text content from the <text> tag and add it to <intro> (ru)
            text_content = text_element.text
            intro_ru.text = text_content
            full_ru.text = text_content
            
            # Translate the content from Russian to English and add it to <intro> (eng)
            translated_text = translator.translate(text_content)
            intro_eng.text = translated_text
            full_eng.text = translated_text

        # Check if the attributes match the specified criteria for <main>
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '20.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'center' and
            text_element.get('leading') == '11.65' and
            text_element.get('indent') == '0.0'
        ):
            # Extract the text content from the <text> tag and add it to <main> (ru)
            text_content = text_element.text
            main_ru.text = text_content
            #if text_content is not None:
             #   full_ru.text += ". " + text_content
            
            # Translate the content from Russian to English and add it to <main> (eng)
            translated_text = translator.translate(text_content)
            main_eng.text = translated_text
            #if translated_text is not None:
              #  full_eng.text += ". " + translated_text

        # Check if the attributes match the specified criteria for <compl>
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '12' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '145.5'
        ):
            # Extract the text content from the <text> tag and add it to <full> (ru)
            text_content = text_element.text
            compl_ru.text = text_content
            #if text_content is not None:
            #    full_ru.text += ". " + text_content

            # Translate the content from Russian to English and add it to <full> (eng)
            translated_text = translator.translate(text_content)
            compl_eng.text = translated_text
            #if translated_text is not None:
            #    full_eng.text += ". " + translated_text
        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '12.349999999999998' and
            text_element.get('indent') == '0.0'
        ):
            p_text_content = text_element.text
            p1_sec_foreword.text = p_text_content
            # Extract the text content from the <text> tag and remove the prefix
            text_content = text_element.text.strip('1 РАЗРАБОТАН')

            # Add the extracted text content to <sdo>
            sdo.text = text_content
        
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '12' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '0.0' and
            'ВНЕСЕН' in text_element.text  # Check if 'ВНЕСЕН' is in the text
        ):
            p_text_content = text_element.text
            p2_sec_foreword.text = p_text_content

            # Extract the text content from the <text> tag
            text_content = text_element.text

            # Extract the text content from the <text> tag and remove the prefix
            text_content = text_content.replace('2 ВНЕСЕН', '').strip()

            # Add the extracted text content to <secretariat>
            secretariat.text = text_content

        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'right' and
            text_element.get('leading') == '11.65'
        ):
            # Extract the text content from the <text> tag
            text_content = text_element.text

            # Add the extracted text content to <ics>
            ics.text = text_content
        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'justify' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '0.0'
        ):
            # Extract the text content from the <text> tag
            text_content = text_element.text

            # Add the extracted text content to <copyright-statement>
            copyright_statement.text = text_content
        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'right' and
            text_element.get('leading') == '11.65' and
            text_element.get('indent') == '0.0'
        ):
            # Extract the text content from the <text> tag
            text_content = text_element.text

            # Add the extracted text content to <copyright-holder>
            copyright_holder.text = text_content
        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '11.65' and
            'УДК' in text_element.text  # Check if 'УДК' is in the text
        ):
            # Extract the text content from the <text> tag
            text_content = text_element.text

            # Check if the text starts with 'УДК '
            if text_content.startswith('УДК '):
                # Remove the 'УДК ' prefix
                udk_value = text_content[len('УДК '):].strip()

                # Set the custom metadata values
                meta_name_udk.text = 'UDK'
                meta_value_udk.text = udk_value
        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '0.3' and
            text_element.text.startswith('Ключевые слова:')
        ):
            # Extract the text content from the <text> tag
            text_content = text_element.text

            # Remove the 'Ключевые слова:' prefix
            keywords = text_content.replace('Ключевые слова:', '').strip()

            # Set the custom metadata values
            meta_name_keywords.text = 'Keywords'
            meta_value_keywords.text = keywords
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '12' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '0.0' and
            'УТВЕРЖДЕН И ВВЕДЕН В ДЕЙСТВИЕ' in text_element.text
        ):
            p_text_content = text_element.text
            p3_sec_foreword.text = p_text_content

            # Extract the date from the text using datetime.strptime
            date_text = text_element.text.split('от ')[-1].split(' №')[0]

            formatted_date  = convert_russian_date(date_text)#datetime.strptime(date_text, '%d %B %Y г.')

            meta_date_nat_meta_ratification.text = formatted_date
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '9.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'right' and
            text_element.get('leading') == '11.65' and
            text_element.get('indent') == '0.0'
        ):
            # Extract the text content
            text_content = text_element.text

            # Remove the prefix 'Дата введения — ' from the text
            date_text = text_content.replace('Дата введения — ', '').strip()

            # Create the <meta-date> element
            meta_date_nat_meta_implementation.text = date_text
        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '11.65' and
            text_element.get('indent') == '25.5' and
            contains_uppercase_word(text_element.text)  # Check for uppercase word
        ):
            # Extract the text content
            text_content = text_element.text


            # Check if the text contains 'ВЗАМЕН' or 'ВПЕРВЫЕ'
            if 'ВЗАМЕН' in text_content or 'ВПЕРВЫЕ' in text_content:
                p4_sec_foreword.text = text_content
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '12.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'center' and
            text_element.get('leading') == '11.65' and
            text_element.get('indent') == '0.0' and
            'Введение' in text_element.text
        ):
            title.text = text_element.text
        elif (
            text_element.get('font') == 'ArialMT' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'false' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '0.0' and
            'Настоящий стандарт в системе' in text_element.text
        ):
            desired_text = text_element.text
            sec_intro_p.text = desired_text
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '9.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'center' and
            text_element.get('leading') == '12.25' and
            text_element.get('indent') == '201.6' 
            ):
            app_line = text_element.get('line_number')

            # Extract the content before and after parentheses
            text_content = text_element.text
            label, annex_type = text_content.split('(')
            label = label.strip()  # Remove leading/trailing spaces

            # Create the <app-group> structure
            app_group = ET.SubElement(back, 'app-group')

            # Create the <app> element
            app = ET.SubElement(app_group, 'app', {'id': 'app_a', 'content-type': 'inform-annex'})

            # Create the <label> element and set its text
            label_element = ET.SubElement(app, 'label')
            label_element.text = label.strip()

            # Create the <annex-type> element and set its text
            annex_type_element = ET.SubElement(app, 'annex-type')
            annex_type_element.text = f'({annex_type.strip()})'
        elif (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'center' and
            text_element.get('leading') == '11.65' and
            text_element.get('indent') == '0.0' and
            'Библиография' in text_element.text
            ):
            bib_line = text_element.get('line_number')

            # Create the <ref-list> structure
            ref_list = ET.SubElement(back, 'ref-list', {'content-type': 'bibl', 'id': 'sec_bibl'})

            # Create the <title> element
            title = ET.SubElement(ref_list, 'title')
            title.text = text_element.text
        
                  
            
    
    for text_element in root.findall('.//text'):
        if (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '10.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            text_element.get('leading') == '11.65' and
            text_element.get('indent') == '0.0'
        ):
            # Extract the text content from the <text> tag
            text_content = text_element.text

            # Use regular expressions to find the pattern with digits and dots
            matches = re.findall(r'(\d+(\.\d+)?)', text_content)
            if matches:
                # Join the matched digits with dots (if there are multiple matches)
                digits_with_dots = '.'.join(match[0] for match in matches)

                # Add the extracted digits with dots to <doc-number>
                doc_number.text = digits_with_dots
                doc_number_nat_meta.text = digits_with_dots

                #std_ref_text = ET.SubElement(std_ref_dated, 'std-ref')
                std_ref_dated.text = text_content
                std_ref_short.text = text_content
                std_ref_undated.text = remove_year_and_em_dash(text_content)
                break


    for link_element in root.findall('.//link'):
        if (
            link_element.get('font') == 'Arial' and
            link_element.get('font_size') == '10.0' and
            link_element.get('bold') == 'false' and
            link_element.get('italic') == 'true' and
            link_element.get('underline') == 'false' and
            link_element.get('alignment') == 'left' and
            link_element.get('leading') == '12.5' and
            link_element.get('indent') == '0.0'
        ):
            # Extract the text content from the <link> element
            link_text = link_element.text

            p_notes_front.text = link_text

            # Extract the URI (www.gost.ru) from the link_text
            uri_start = link_text.find('www.')
            if uri_start != -1:
                uri = link_text[uri_start:]
            else:
                uri = ""

            # Add the URI as <uri>
            uri_element = ET.SubElement(p_notes_front, 'uri')
            uri_element.text = uri

    line_number_list = [] 

    # Find all <text> elements with the specified attributes
    section_texts = []
    for text_element in root.findall('.//text'):
        if (
            text_element.get('font') == 'Arial' and
            text_element.get('font_size') == '12.0' and
            text_element.get('bold') == 'true' and
            text_element.get('italic') == 'false' and
            text_element.get('underline') == 'false' and
            text_element.get('alignment') == 'left' and
            (
                text_element.get('leading') == '11.65' or
                text_element.get('leading') == '11.75'
            ) and
            text_element.get('indent') is not None #== '25.5'
        ):
            section_texts.append(text_element.text)
            # Get the line number attribute
            line_number = text_element.get('line_number')
            line_number_list.append(line_number)

    largest_line = find_largest_line_number(xml_file)   


    if app_line is not None:
        line_number_list.append(app_line)
    if bib_line is not None:    
        line_number_list.append(bib_line)

    line_number_list.append(largest_line)
    print(line_number_list)
    numbers = [int(value) for value in line_number_list]
    print(numbers)


    # Create <sec> elements based on the section_texts
    for i, section_text in enumerate(section_texts, start=1):
        sec = ET.SubElement(body, 'sec', {'id': f'sec_{i}'})

        # Create <label> element with the section number
        label = ET.SubElement(sec, 'label')
        label.text = str(i)

        # Create <title> element with the section name
        title = ET.SubElement(sec, 'title')

        # Use regular expressions to remove leading digits
        section_text = re.sub(r'^\d+\s*', '', section_text)  
        title.text = section_text

    # Define regular expression patterns for matching id attributes
    sec_id_pattern = re.compile(r'sec_\d+')
    app_id_pattern = re.compile(r'app_[a-zA-Z]+')

    # Define a function to recursively search for matching tags
    def find_tags_by_id_pattern(element, pattern):
        matching_tags = []

        if 'id' in element.attrib and pattern.match(element.attrib['id']):
            matching_tags.append(element)

        for child in element:
            matching_tags.extend(find_tags_by_id_pattern(child, pattern))

        return matching_tags

    # Search for matching <sec> tags
    matching_sec_tags = find_tags_by_id_pattern(standard_root, sec_id_pattern)

    # Search for matching <app> tags
    matching_app_tags = find_tags_by_id_pattern(standard_root, app_id_pattern)

    # Search for matching <ref-list> tags with content-type="bibl" and id="sec_bibl"
    matching_ref_list_tags = standard_root.findall(".//ref-list[@content-type='bibl'][@id='sec_bibl']")

    matching_ref_list_ids = [ref_list_tag.attrib['id'] for ref_list_tag in matching_ref_list_tags]

    # Collect all the collected ids into one list
    all_ids = [tag.attrib['id'] for tag in matching_sec_tags + matching_app_tags] + matching_ref_list_ids

    # Initialize an empty dictionary to store the line numbers for each id
    line_numbers_dist = {}

    # Pair adjacent numbers and create the mapping
    for i in range(len(all_ids)):
        if i < len(all_ids) - 1:
            start_line = numbers[i] + 1
            end_line = numbers[i + 1] - 1
        else:
            # For the last id, use the last number in the list as the end line
            start_line = numbers[i] + 1
            end_line = numbers[-1]

        # Create a list of line numbers between start_line and end_line (inclusive)
        line_numbers = list(range(start_line, end_line + 1))

        # Add the id and its associated line numbers to the dictionary
        line_numbers_dist[all_ids[i]] = line_numbers

    # Print the resulting mapping
    for key, value in line_numbers_dist.items():
        print(f"{key}: {value}")

    def parse_xml_by_line_number(xml_file, target_line_number):
        try:
            tree = ET.parse(xml_file)
            root = tree.getroot()
        except ET.ParseError:
            print("Error parsing XML file.")
            return []

        parsed_texts = []


        for element in root.iter():
            if element.tag in ('text', 'list_item'):
                line_number = int(element.get("line_number", 0))

                if line_number == target_line_number:
                    text = ''
                    if element.text is not None:
                        text += element.text
                    if element.tail is not None:
                        text += element.tail

                    parsed_text = text.strip()
                    if parsed_text:
                        parsed_texts.append(parsed_text)

        return parsed_texts
    
    def replace_brackets_with_tags(text):
        def replace(match):
            index = match.group(1)
            xref_element = ET.Element('xref')
            xref_element.set('ref-type', 'bibr')
            xref_element.set('rid', f'biblref_{index}')
            xref_element.text = f'[{index}]'
            return ET.tostring(xref_element).decode()

        pattern = r'\[(\d+)\]'
        replaced_text = re.sub(pattern, replace, text)
        return replaced_text

    def find_section_id_with_title(tree, title_keywords):
        for section in tree.findall(".//sec"):
            title = section.find(".//title")
            if title is not None and any(keyword in (title.text or "").lower() for keyword in title_keywords):
                section_id = section.get("id")
                return section_id
        return None

    xml_tree_dummy = ET.ElementTree(standard_root)  # Replace with your XML tree
    
    keywords_to_find_terms = ['термины', 'определения']
    matching_section_id_terms = find_section_id_with_title(xml_tree_dummy, keywords_to_find_terms)
    
    keywords_to_find_refs = ['ссылки']
    matching_section_id_refs = find_section_id_with_title(xml_tree_dummy, keywords_to_find_refs)
      
    pattern_id = r'^(\d+(\.\d+)*)'
    
    for sec in range(len(list(line_numbers_dist.values()))):
        tag_id = standard_root.findall(f".//*[@id='{list(line_numbers_dist.keys())[sec]}']")
        
        for line_number in list(line_numbers_dist.values())[sec]: 
            parsed_text = parse_xml_by_line_number(xml_file, line_number)
            
            if parsed_text is not None and len(parsed_text) == 1:
                match = re.match(pattern_id, parsed_text[0])
                
                if 'П р и м е ч а н и е' in parsed_text[0]:
                    non_note = ET.SubElement(tag_id[0], 'non-normative-note')
                    label_note = ET.SubElement(non_note, 'label')
                    label_note.text = 'Примечание — '
                    text_note = ET.SubElement(non_note, 'p')
                    text_note.text = parsed_text[0].replace('П р и м е ч а н и е', '').strip()
                
                elif list(line_numbers_dist.keys())[sec] == matching_section_id_refs:                    
                    if 'В настоящем стандарте' in parsed_text[0]:
                        # Add parsed_text containing 'В настоящем стандарте' as a <p> tag
                        p_tag = ET.SubElement(tag_id[0], 'p')
                        p_tag.text = parsed_text[0]                    
                    elif 'В настоящем стандарте' not in parsed_text[0]:
                        #description_gost = ""
                
                        if parsed_text[0].startswith('ГОСТ'):                              
                            # Regular expression pattern to match ГОСТ references with variations
                            gost_pattern = r'ГОСТ\s+(?:ISO/IEC|Р|Р\s+ИСО)?\s+\d[\d./-]*'
                            matches_gost = re.findall(gost_pattern, parsed_text[0])

                            if matches_gost:
                                ref_gost = matches_gost[0]#gost_references.extend(matches)
                                description_gost = re.sub(gost_pattern, '', parsed_text[0]).strip()

                                ref_list_tag = ET.SubElement(tag_id[0], 'ref-list')
                                ref_list_tag.set('content-type', 'normref')
                                ref_list_tag_id = ET.SubElement(ref_list_tag, 'ref')
                                ref_list_tag_id.set('id', 'normref')
                                ref_list_tag_std = ET.SubElement(ref_list_tag_id, 'std')
                                ref_list_tag_std.set('type', 'undated')
                                ref_list_tag_std.set('std-id', 'ГОСТ')

                                #if ref_gost or (name_gost and name_gost_continue):
                                ref_list_tag_std_ref = ET.SubElement(ref_list_tag_std, 'std-ref')
                                ref_list_tag_std_ref.text = ref_gost
                                ref_list_tag_std_title = ET.SubElement(ref_list_tag_std, 'title')
                                ref_list_tag_std_title.text = description_gost# + name_gost_continue
                        #elif parsed_text[0][0].islower():  # Check if the first character is lowercase
                        #    ref_list_tag_std_p = ET.SubElement(ref_list_tag_std, 'p')
                         #   ref_list_tag_std_p.text = parsed_text[0]
                            #description_gost += parsed_text[0]
                            
                        #ref_list_tag_std_title = ET.SubElement(ref_list_tag_std, 'title')
                        #ref_list_tag_std_title.text = description_gost
                elif list(line_numbers_dist.keys())[sec] == matching_section_id_terms:    
                    term_pattern = r'^\d+(\.\d+)+'
                    term_id = re.match(term_pattern, parsed_text[0])
                    if term_id and 'В настоящем стандарте' not in parsed_text[0]:
                        # Extract the label and definition
                        # Ensure that there's a ':' character in the parsed text before splitting
                        if ':' in parsed_text[0]:
                            label, definition = parsed_text[0].split(': ', 1)
                            label = label.strip()  # Extracted label
                            definition = definition.strip()  # Extracted definition
                        else:
                            # Handle the case where ':' is not present in parsed_text[0]
                            label = parsed_text[0].strip()  # Extracted label (full text)
                            definition = None  # None if there is no definition

                        # Create the term-sec structure
                        term_sec_tag = ET.SubElement(tag_id[0], f'term-sec')
                        term_sec_tag.set('id', f'term_sec_{term_id[0]}')
                        
                        # Add definition if it is not None and not empty
                        term_sec_label_tag = ET.SubElement(term_sec_tag, 'label')
                        term_sec_label_tag.text = term_id[0]
                        
                        term_entry_tag = ET.SubElement(term_sec_tag, 'tbx:termEntry')
                        term_entry_tag.set('id', f'term_entry_{term_id[0]}')
                        term_entry_id = ET.SubElement(term_entry_tag, 'tbx:langSet')
                        term_entry_id.set('xml:lang', 'ru')
                        
                        # Create a set to keep track of processed term_sec_tags
                        #processed_term_sec_tags = set()

                        #for index in range(len(text_term_table_list_1)):
                            # Set the display attribute to "frame" if definition is None or empty
                        if definition is None or definition == "":
                            term_sec_tag.set('display', 'frame')
                            definition_frame_tag = ET.SubElement(term_entry_id, 'definition')
                            source_frame_tag = ET.SubElement(term_entry_id, 'source')
            
                        if definition is not None and definition != "":
                            definition_tag = ET.SubElement(term_entry_id, 'tbx:definition')
                            definition_tag.text = definition

                        tig_tag = ET.SubElement(term_entry_id, 'tbx:tig')
                        term_tag = ET.SubElement(tig_tag, 'tbx:term')
                        partOfSpeech_tag = ET.SubElement(tig_tag, 'tbx:partOfSpeech')
                        partOfSpeech_tag.set('value', "noun")
                        term_tag.text = re.sub(r'^\d+(\.\d+)+\s', '', label)
                    
                    elif 'В настоящем стандарте' in parsed_text[0]:
                        # Add parsed_text containing 'В настоящем стандарте' as a <p> tag
                        #p_id = match.group(1)
                        p_tag = ET.SubElement(tag_id[0], 'p')
                        p_tag.text = parsed_text[0]
                        #p_id_tag.set('id', f'p_{term_id}')


                elif list(line_numbers_dist.keys())[sec].startswith('app_'):
                    # Check if the line starts with a format like 'А.1', 'А.2', etc.
                    app_pattern = r'^(?:[А-Я]\.\d+)'  # Modify this pattern as needed
                    match_app = re.match(app_pattern, parsed_text[0])
                    if match_app:
                        app_id = match_app.group(0).replace('.', '').lower()  # Convert to lowercase and remove '.'
                        p_tag = ET.SubElement(tag_id[0], 'p')
                        p_tag.set('id', f'p_{app_id}')
                        p_tag.text = parsed_text[0]            
                elif list(line_numbers_dist.keys())[sec] == 'sec_bibl':
                    # Check if the line contains any digit between [ and ]
                    cit_pattern = r'\[(\d+)\]'
                    match_cit = re.search(cit_pattern, parsed_text[0])
                    if match_cit:
                        cit_id = match_cit.group(0) 
                        cit_tag = ET.SubElement(tag_id[0], 'ref')
                        digit_id = re.sub(cit_pattern, r'\1', cit_id)
                        cit_tag.set('id', f"biblref_{digit_id}") #cit_id}')
                        label_cit = ET.SubElement(cit_tag, 'label')
                        label_cit.text = match_cit.group(0)
                        cit_el = ET.SubElement(cit_tag, 'element-citation')
                        source = ET.SubElement(cit_el, 'source')
                        source.text = re.sub(cit_pattern, '', parsed_text[0])
                elif match:
                    p_id = match.group(1)
                    p_id_tag = ET.SubElement(tag_id[0], 'p')
                    p_id_tag.set('id', f'p_{p_id}')
                    p_id_tag.text = replace_brackets_with_tags(parsed_text[0])
                else:
                    p_tag = ET.SubElement(tag_id[0], 'p')
                    p_tag.text = replace_brackets_with_tags(parsed_text[0])
    
    text_term_table_list_1, text_term_table_list_2 = [], []
    for t in range(len(root.findall(".//text_term_table_1"))):
        text_term_table_list_1_t = root.findall(".//text_term_table_1")[t].text
        text_term_table_list_2_t = root.findall(".//text_term_table_2")[t].text
        text_term_table_list_1.append(text_term_table_list_1_t)
        text_term_table_list_2.append(text_term_table_list_2_t)
    
    print(text_term_table_list_1)
    print(text_term_table_list_2)
    
    source_tags = standard_root.findall('.//source')#(".//{urn:iso:std:iso:30042:ed-1}source")
    def_tags = standard_root.findall(".//definition")
    
    #for s in range(len(source_tags)):
    #    source_tags[s].text = text_term_table_list_2[s]
    
    print(source_tags)
    print(def_tags)

    # Create an iterator that cycles through text_term_table_list_1 indefinitely
    text_iterator_1 = itertools.cycle(text_term_table_list_1)
    text_iterator_2 = itertools.cycle(text_term_table_list_2)
    
    # Iterate through definition tags and assign text from the list
    for def_tag in def_tags:
        if def_tag.text is None:
            def_tag.text = next(text_iterator_1)
    for source_tag in source_tags:
        if source_tag.text is None:
            source_tag.text = replace_brackets_with_tags(next(text_iterator_2))
    
    intro_text = standard_root.findall('.//intro')[0].text
    main_text = standard_root.findall('.//main')[0].text
    compl_text = standard_root.findall('.//intro')[0].text
    full_ru_tag, full_en_tag = standard_root.findall('.//full')[0].text, standard_root.findall('.//full')[1].text
    
    if (intro_text and main_text and compl_text) is not None:
        full_ru_tag.text = intro_text + main_text + compl_text
        full_en_tag.text = Translator(full_ru_tag.text)
    
    # Create a new XML tree with the modified structure
    new_tree = ET.ElementTree(standard_root)

    # Save the new XML to a file
    new_tree.write(xml_structure, encoding='utf-8', xml_declaration=True)
    #return standard_root#, xml_tree_dummy


class PDFXMLViewer(QMainWindow):
    def __init__(self):
        super().__init__()

        self.initUI()

    def initUI(self):
        self.setWindowTitle("PDF and XML Viewer")
        self.setGeometry(100, 100, 800, 600)

        # Create a central widget and layout for the main window
        central_widget = QWidget(self)
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # Create a splitter to divide the window into left and right sections
        splitter = QSplitter(Qt.Horizontal)

        # Create a PDF viewer widget on the left side within a scroll area
        pdf_scroll_area = QScrollArea()
        self.pdf_viewer = QWidget()
        self.pdf_viewer.setLayout(QVBoxLayout())  # Create a layout for the PDF viewer
        pdf_scroll_area.setWidget(self.pdf_viewer)
        pdf_scroll_area.setWidgetResizable(True)  # Make the scroll area resizable
        splitter.addWidget(pdf_scroll_area)

        # Create an XML viewer widget on the right side
        self.xml_viewer = QTextEdit(self)
        splitter.addWidget(self.xml_viewer)

        layout.addWidget(splitter)

        # Create a button to open the folder containing PDF and XML files
        open_button = QPushButton("Open PDF File", self)
        open_button.clicked.connect(self.openPDF)
        layout.addWidget(open_button)

    def openPDF(self):
        pdf_file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File")

        if pdf_file_path:
            # Display the selected PDF file
            self.displayPDF(pdf_file_path)

            # Generate XML from the selected PDF and display it
            self.generateAndDisplayXML(pdf_file_path)

    def displayPDF(self, pdf_path):
        # Clear any existing widgets in the PDF viewer area
        pdf_layout = self.pdf_viewer.layout()
        while pdf_layout.count():
            item = pdf_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()

        # Use PyMuPDF to render and display all pages of the PDF file within the scrollable area
        pdf_document = fitz.open(pdf_path)

        for page_num in range(len(pdf_document)):
            pdf_page = pdf_document.load_page(page_num)
            pixmap = pdf_page.get_pixmap()
            qt_pixmap = QPixmap.fromImage(QImage(pixmap.samples, pixmap.width, pixmap.height, pixmap.stride, QImage.Format_RGB888))
            label = QLabel(self)
            label.setPixmap(qt_pixmap)
            pdf_layout.addWidget(label)

    def generateAndDisplayXML(self, pdf_path):
        # Generate XML from the PDF file
        base_filename = os.path.splitext(os.path.basename(pdf_path))[0]
        output_folder = os.path.join(os.path.dirname(pdf_path), base_filename)

        # Ensure the output folder exists
        os.makedirs(output_folder, exist_ok=True)

        # Convert PDF to DOCX while preserving styles
        docx_file = os.path.join(output_folder, f"{base_filename}.docx")
        pdf_to_docx(pdf_path, docx_file)

        # Create a subfolder for images from the DOCX file
        image_subfolder = os.path.join(output_folder, f"output_images_{base_filename}")
        os.makedirs(image_subfolder, exist_ok=True)

        # Helper function to convert DOCX paragraph styles to XML element attributes
        defaultXML = os.path.join(output_folder, f"{base_filename}_default.xml")
        docx_to_xml(docx_file, defaultXML)
        
        xml = os.path.join(output_folder, f"{base_filename}.xml")
        defaultXML2XML(defaultXML, xml)

        # Function to save images from DOCX to the subfolder
        save_images_from_docx(docx_file, image_subfolder)

        # Display the generated XML in the XML viewer
        xml_file = os.path.join(output_folder, f"{base_filename}.xml")
        with open(xml_file, 'r', encoding='utf-8') as xml_file:
            xml_content = xml_file.read()
            #self.xml_viewer.setPlainText(xml_content)
            xml_lines = xml_content.splitlines()
            xml_cursor = self.xml_viewer.textCursor()
            xml_cursor.movePosition(QTextCursor.Start)
            
            # Define more visible colors for tags on a white background
            tag_colors = [
                QColor("#E60000"),  # Red
                QColor("#339933"),  # Green
                QColor("#0000E6"),  # Blue
                QColor("#9933CC"),  # Purple
            ]

            for line in xml_lines:
                # Split the line into tags and text
                tags_and_text = line.split('<')
                
                for tag_and_text in tags_and_text:
                    # Split tags and text into individual tags
                    individual_tags = tag_and_text.split('>')
                    for tag in individual_tags[:-1]:
                        # Get the tag name (without '<' and '/') and set the tag's color
                        tag_name = tag.split('/')[0]
                        color_index = hash(tag_name) % len(tag_colors)
                        tag_color = tag_colors[color_index]
                        
                        # Add the tag with color to the QTextEdit
                        format = QTextCharFormat()
                        format.setForeground(tag_color)
                        xml_cursor.setCharFormat(format)
                        xml_cursor.insertText('<' + tag_name + '>')
                        
                        # Move to the next position
                        xml_cursor.movePosition(QTextCursor.End)

                    # Add the last part (text) of the split line
                    xml_cursor.insertText(individual_tags[-1])
                    # Add a newline to separate XML lines
                    xml_cursor.insertText('\n')
                
            # Scroll to the beginning of the document
            self.xml_viewer.moveCursor(QTextCursor.Start)

def main():
    app = QApplication(sys.argv)
    viewer = PDFXMLViewer()
    viewer.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
